import os
import docx
from docx.shared import Inches
class doc:
    def __init__(self, filesize, filename,generator,WORD_SEPARATOR):
        self.filesize = filesize
        self.filename = filename
        self.generator = generator
        self.word_separator = WORD_SEPARATOR

    def execute(self):
        for i in range(0,len(self.filename)):
            statinfo = 0
            file_size = int(self.filesize[i])
            file_name = self.filename[i]
            doc = docx.Document()
            doc.save(file_name)
            while statinfo < file_size:
                #f = open(file_name,"a")
                doc = docx.Document(file_name)
                temptxt = ""
                temptxt = self.generator.generate(self.word_separator)
                #print(temptxt)
                doc.add_heading('Heading, level 1', level=1)
                for i in range(0, 1000):
                    #f.write(self.generator.generate(self.word_separator))
                    temptxt = temptxt + " " + self.generator.generate(self.word_separator)
                try:
                    doc.add_paragraph(temptxt)
                except ValueError as e:
                    print("ATTENTION:"+temptxt)
                doc.add_picture('1.jpg')
                doc.add_paragraph('first item in unordered list', style='ListBullet')
                doc.add_paragraph('first item in ordered list', style='ListNumber')
                #f.close()
                doc.save(file_name)
                statinfo = os.stat(file_name).st_size
                print(statinfo)